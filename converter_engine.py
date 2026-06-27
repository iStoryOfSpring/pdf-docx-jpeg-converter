"""
PDF·Word·JPG 互转引擎
武汉纺织大学管理学院媒体运营部

提供 6 条转换路径：
  PDF → Word      (pdf2docx)
  PDF → JPEG      (PyMuPDF 光栅化)
  Word → PDF      (COM — 兼容 MS Word / WPS Office)
  Word → JPEG     (Word→PDF→JPEG 链式，COM)
  JPEG → PDF      (PyMuPDF 嵌入)
  JPEG → Word     (python-docx 插入图片)
  .doc → .docx    (COM 桥接 — 兼容 MS Word / WPS Office)
"""

import os
import tempfile
from typing import Callable, Optional


class ConverterEngine:
    """文档格式互转核心引擎。"""

    SUPPORTED_EXT = (".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png")
    FORMAT_MAP = {
        ".pdf": "PDF",
        ".docx": "Word",
        ".doc": "Word",
        ".jpg": "JPEG",
        ".jpeg": "JPEG",
        ".png": "PNG",
    }

    # ── 静态工具方法 ──────────────────────────────────────────

    @staticmethod
    def detect_format(path: str) -> str:
        """根据扩展名返回格式标签：'PDF' / 'Word' / 'JPEG' / 'PNG'"""
        ext = os.path.splitext(path)[1].lower()
        if ext not in ConverterEngine.FORMAT_MAP:
            raise ValueError(f"不支持的文件格式：{ext}")
        return ConverterEngine.FORMAT_MAP[ext]

    @staticmethod
    def _resolve_conflict(path: str, strategy: str = "rename") -> Optional[str]:
        """
        根据冲突策略处理输出路径。

        参数:
            path:     期望的输出路径
            strategy: 'rename' | 'overwrite' | 'skip'

        返回:
            最终路径（skip 且存在时返回 None）
        """
        if not os.path.exists(path):
            return path

        if strategy == "overwrite":
            return path

        if strategy == "skip":
            return None

        # strategy == 'rename' → 追加 _1, _2 ...
        base, ext = os.path.splitext(path)
        counter = 1
        while True:
            new_path = f"{base}_{counter}{ext}"
            if not os.path.exists(new_path):
                return new_path
            counter += 1

    @staticmethod
    def _format_size(size_bytes: int) -> str:
        """将字节数转成人类可读字符串。"""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.1f} MB"

    # ── COM 办公套件探测（MS Word / WPS）───────────────────

    # WPS 文字组件的 COM ProgID（不同版本可能不同）
    _WPS_WORD_PROGIDS = [
        "Word.Application",      # Microsoft Word
        "WPS.Application",       # WPS Office 文字（新版）
        "KWPS.Application",      # WPS Office 文字（旧版 / 金山 WPS）
        "WPS.WpsApplication",    # 某些 WPS 版本
    ]

    @staticmethod
    def _create_com_app():
        """
        按优先级依次尝试 MS Word → WPS → KWPS，创建并返回办公套件 COM 对象。

        Returns:
            (app_object, progid)

        Raises:
            RuntimeError: 如果所有 COM ProgID 均不可用
        """
        import win32com.client as wc

        last_error = None
        for progid in ConverterEngine._WPS_WORD_PROGIDS:
            try:
                app = wc.Dispatch(progid)
                app.Visible = False
                app.DisplayAlerts = 0
                return app, progid
            except Exception as e:
                last_error = e
                continue

        raise RuntimeError(
            "未检测到可用的办公套件。请安装 Microsoft Word 或 WPS Office。\n"
            f"尝试过的组件：{', '.join(ConverterEngine._WPS_WORD_PROGIDS)}\n"
            f"最后错误：{last_error}"
        )

    @staticmethod
    def _docx_to_pdf_via_com(docx_path: str, output_pdf_path: str, app) -> None:
        """
        通过 COM 将 .docx 另存为 PDF。

        支持 MS Word 和 WPS Office — 两者 API 高度兼容。

        Args:
            app: 已初始化的 COM app 对象
        """
        doc = None
        try:
            doc = app.Documents.Open(docx_path, ReadOnly=True)
            # 17 = wdFormatPDF (MS Word)
            # WPS 同样支持 FileFormat=17 导出 PDF
            doc.SaveAs2(output_pdf_path, FileFormat=17)
        finally:
            if doc is not None:
                doc.Close(SaveChanges=False)

    # ── .doc → .docx 桥接（COM）────────────────────────────

    @staticmethod
    def _doc_to_docx_via_com(doc_path: str, app, progress_callback: Optional[Callable] = None) -> str:
        """
        通过 COM 将 .doc 另存为临时 .docx。

        Args:
            doc_path: .doc 文件路径
            app:      已初始化的 COM app（MS Word 或 WPS）

        Returns:
            临时 .docx 文件的路径
        """
        doc = None
        try:
            if progress_callback:
                progress_callback(5, "正在打开 .doc 文件...")

            doc = app.Documents.Open(doc_path, ReadOnly=True)

            if progress_callback:
                progress_callback(40, "正在另存为 .docx...")

            fd, tmp_docx = tempfile.mkstemp(suffix=".docx", prefix="cnv_doc_")
            os.close(fd)

            # 16 = wdFormatXMLDocument (.docx)
            doc.SaveAs2(tmp_docx, FileFormat=16)

            if progress_callback:
                progress_callback(90, "临时 .docx 生成完毕")

            return tmp_docx
        finally:
            if doc is not None:
                doc.Close(SaveChanges=False)

    @staticmethod
    def _ensure_docx(path: str, app, progress_callback: Optional[Callable] = None) -> tuple:
        """
        若输入是 .doc 则先转为临时 .docx。

        Args:
            app: 已初始化的 COM app（MS Word 或 WPS）

        Returns:
            (actual_input_path, tmp_docx_path_to_cleanup_or_None)
        """
        ext = os.path.splitext(path)[1].lower()
        if ext != ".doc":
            return path, None  # 无需转换

        tmp_docx = ConverterEngine._doc_to_docx_via_com(path, app, progress_callback)
        return tmp_docx, tmp_docx

    # ── 6 条转换路径 ───────────────────────────────────────────

    # 1. PDF → Word
    def pdf_to_word(
        self,
        pdf_path: str,
        output_dir: str,
        conflict: str = "rename",
        progress_callback: Optional[Callable] = None,
    ) -> dict:
        """
        将 PDF 转换为 Word (docx)。

        返回: {'success': bool, 'output_files': list[str], 'error': str|None}
        """
        from pdf2docx import Converter as PDF2DocxConverter

        try:
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            out_path = os.path.join(output_dir, f"{base_name}.docx")
            out_path = self._resolve_conflict(out_path, conflict)
            if out_path is None:
                return {"success": False, "output_files": [], "error": "目标文件已存在，冲突策略为'跳过'"}

            if progress_callback:
                progress_callback(10, "正在解析 PDF 结构...")

            cv = PDF2DocxConverter(pdf_path)
            cv.convert(out_path)
            cv.close()

            if progress_callback:
                progress_callback(100, "转换完成")

            return {"success": True, "output_files": [out_path], "error": None}
        except Exception as e:
            return {"success": False, "output_files": [], "error": f"PDF→Word 转换失败：{e}"}

    # 2. PDF → JPEG
    def pdf_to_jpeg(
        self,
        pdf_path: str,
        output_dir: str,
        dpi: int = 200,
        quality: int = 90,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None,
        conflict: str = "rename",
        progress_callback: Optional[Callable] = None,
    ) -> dict:
        """
        将 PDF 页面逐页光栅化为 JPEG。

        start_page / end_page 为 1-based；None 表示全部。
        """
        import fitz  # PyMuPDF

        try:
            doc = fitz.open(pdf_path)
            total_pages = len(doc)

            # 规范化页码范围
            first = 1
            last = total_pages
            if start_page is not None:
                first = max(1, start_page)
            if end_page is not None:
                last = min(total_pages, end_page)
            if first > last:
                doc.close()
                return {"success": False, "output_files": [], "error": f"页码范围无效（{first}-{last}）"}

            page_count = last - first + 1
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            output_files = []

            if progress_callback:
                progress_callback(0, f"正在渲染 {page_count} 页...")

            for i, page_num in enumerate(range(first, last + 1)):
                page = doc.load_page(page_num - 1)  # fitz 是 0-based
                pix = page.get_pixmap(dpi=dpi)

                if page_count == 1:
                    out_name = f"{base_name}.jpg"
                else:
                    out_name = f"{base_name}_page_{page_num:03d}.jpg"

                out_path = os.path.join(output_dir, out_name)
                out_path = self._resolve_conflict(out_path, conflict)
                if out_path is None:
                    continue  # skip

                pix.save(out_path, jpg_quality=quality)
                output_files.append(out_path)

                if progress_callback:
                    percent = int((i + 1) / page_count * 100)
                    progress_callback(percent, f"正在渲染第 {page_num}/{last} 页...")

            doc.close()

            if not output_files:
                return {"success": False, "output_files": [], "error": "没有生成任何文件（全部跳过或无法渲染）"}

            if progress_callback:
                progress_callback(100, "渲染完成")

            return {"success": True, "output_files": output_files, "error": None}
        except Exception as e:
            return {"success": False, "output_files": [], "error": f"PDF→JPEG 转换失败：{e}"}

    # 3. Word → PDF
    def word_to_pdf(
        self,
        docx_path: str,
        output_dir: str,
        conflict: str = "rename",
        original_name: Optional[str] = None,
        app=None,
        progress_callback: Optional[Callable] = None,
    ) -> dict:
        """
        将 Word 文档转换为 PDF。

        通过 COM（MS Word 或 WPS）实现，保真度最高。

        Args:
            original_name: 如传入则为输出文件基名（用于 .doc 桥接场景）
            app:           可复用的 COM app 对象（避免重复创建/销毁）
        """
        own_app = False
        try:
            base_name = original_name or os.path.splitext(os.path.basename(docx_path))[0]
            out_path = os.path.join(output_dir, f"{base_name}.pdf")
            out_path = self._resolve_conflict(out_path, conflict)
            if out_path is None:
                return {"success": False, "output_files": [], "error": "目标文件已存在，冲突策略为'跳过'"}

            if progress_callback:
                progress_callback(20, "正在启动办公套件...")

            if app is None:
                app, _ = self._create_com_app()
                own_app = True

            self._docx_to_pdf_via_com(docx_path, out_path, app)

            if progress_callback:
                progress_callback(100, "转换完成")

            return {"success": True, "output_files": [out_path], "error": None}
        except Exception as e:
            return {"success": False, "output_files": [], "error": f"Word→PDF 转换失败：{e}"}
        finally:
            if own_app and app is not None:
                try:
                    app.Quit()
                except Exception:
                    pass

    # 4. Word → JPEG （链式：Word→临时PDF→JPEG）
    def word_to_jpeg(
        self,
        docx_path: str,
        output_dir: str,
        dpi: int = 200,
        quality: int = 90,
        conflict: str = "rename",
        original_name: Optional[str] = None,
        app=None,
        progress_callback: Optional[Callable] = None,
    ) -> dict:
        """
        将 Word 文档转换为 JPEG 图片。

        步骤：(1) docx → 临时 PDF（COM） (2) 临时 PDF → JPEG（PyMuPDF）

        Args:
            original_name: 如传入则为输出文件基名（用于 .doc 桥接场景）
            app:           可复用的 COM app 对象（避免重复创建/销毁）
        """
        tmp_pdf = None
        own_app = False
        try:
            # Step 1: Word → 临时 PDF
            if progress_callback:
                progress_callback(10, "正在生成中间 PDF...")

            fd, tmp_pdf = tempfile.mkstemp(suffix=".pdf", prefix="cnv_")
            os.close(fd)

            if app is None:
                app, _ = self._create_com_app()
                own_app = True

            self._docx_to_pdf_via_com(docx_path, tmp_pdf, app)

            # Step 2: 临时 PDF → JPEG
            if progress_callback:
                progress_callback(40, "正在渲染页面...")

            result = self.pdf_to_jpeg(
                tmp_pdf,
                output_dir,
                dpi=dpi,
                quality=quality,
                conflict=conflict,
                # 传递一个包装回调，把进度映射到 40%-100%
                progress_callback=(
                    lambda pct, msg: progress_callback(40 + int(pct * 0.6), msg)
                    if progress_callback
                    else None
                ),
            )

            # 将输出文件名从临时 PDF 名重命名为原始文件名
            if result["success"]:
                target_base = original_name or os.path.splitext(os.path.basename(docx_path))[0]
                renamed_files = []
                for old_path in result["output_files"]:
                    old_name = os.path.basename(old_path)
                    tmp_base = os.path.splitext(os.path.basename(tmp_pdf))[0]
                    new_name = old_name.replace(tmp_base, target_base)
                    new_path = os.path.join(output_dir, new_name)
                    new_path = self._resolve_conflict(new_path, conflict)
                    if new_path is None:
                        try:
                            os.unlink(old_path)
                        except OSError:
                            pass
                        continue
                    os.rename(old_path, new_path)
                    renamed_files.append(new_path)
                result["output_files"] = renamed_files
                if not renamed_files:
                    result["success"] = False
                    result["error"] = "没有生成任何文件（全部跳过）"

            if progress_callback:
                progress_callback(100, "转换完成")

            return result
        except Exception as e:
            return {"success": False, "output_files": [], "error": f"Word→JPEG 转换失败：{e}"}
        finally:
            if own_app and app is not None:
                try:
                    app.Quit()
                except Exception:
                    pass
            if tmp_pdf and os.path.exists(tmp_pdf):
                try:
                    os.unlink(tmp_pdf)
                except OSError:
                    pass

    # 5. JPEG → PDF
    def jpeg_to_pdf(
        self,
        img_path: str,
        output_dir: str,
        conflict: str = "rename",
        progress_callback: Optional[Callable] = None,
    ) -> dict:
        """
        将图片嵌入为单页 PDF。
        """
        import fitz  # PyMuPDF

        try:
            base_name = os.path.splitext(os.path.basename(img_path))[0]
            out_path = os.path.join(output_dir, f"{base_name}.pdf")
            out_path = self._resolve_conflict(out_path, conflict)
            if out_path is None:
                return {"success": False, "output_files": [], "error": "目标文件已存在，冲突策略为'跳过'"}

            if progress_callback:
                progress_callback(30, "正在读取图片...")

            # 读取图片获取尺寸
            from PIL import Image

            img = Image.open(img_path)
            img_w, img_h = img.size
            img_dpi = img.info.get("dpi", (72, 72))
            # PIL dpi 是 (x_dpi, y_dpi) 元组
            dpi_x = img_dpi[0] if isinstance(img_dpi, (tuple, list)) else img_dpi
            dpi_y = img_dpi[1] if isinstance(img_dpi, (tuple, list)) else img_dpi

            if progress_callback:
                progress_callback(60, "正在生成 PDF...")

            # 创建 PDF（尺寸单位：point = 1/72 inch）
            # 根据 dpi 和像素计算 point 尺寸
            page_w = img_w / dpi_x * 72
            page_h = img_h / dpi_y * 72

            doc = fitz.open()
            page = doc.new_page(width=page_w, height=page_h)

            with open(img_path, "rb") as f:
                img_bytes = f.read()

            rect = fitz.Rect(0, 0, page_w, page_h)
            page.insert_image(rect, stream=img_bytes)
            doc.save(out_path)
            doc.close()
            img.close()

            if progress_callback:
                progress_callback(100, "转换完成")

            return {"success": True, "output_files": [out_path], "error": None}
        except Exception as e:
            return {"success": False, "output_files": [], "error": f"JPEG→PDF 转换失败：{e}"}

    # 6. JPEG → Word
    def jpeg_to_word(
        self,
        img_path: str,
        output_dir: str,
        conflict: str = "rename",
        progress_callback: Optional[Callable] = None,
    ) -> dict:
        """
        将图片插入到 Word 文档中。
        """
        from docx import Document
        from docx.shared import Inches, Cm
        from PIL import Image

        try:
            base_name = os.path.splitext(os.path.basename(img_path))[0]
            out_path = os.path.join(output_dir, f"{base_name}.docx")
            out_path = self._resolve_conflict(out_path, conflict)
            if out_path is None:
                return {"success": False, "output_files": [], "error": "目标文件已存在，冲突策略为'跳过'"}

            if progress_callback:
                progress_callback(30, "正在读取图片尺寸...")

            # 获取图片宽高以适配页面
            img = Image.open(img_path)
            img_w, img_h = img.size

            # 获取图片 DPI 用于计算实际尺寸
            img_dpi = img.info.get("dpi", (96, 96))
            dpi_x = img_dpi[0] if isinstance(img_dpi, (tuple, list)) else img_dpi
            dpi_y = img_dpi[1] if isinstance(img_dpi, (tuple, list)) else img_dpi

            # 转换为英寸
            width_inch = img_w / dpi_x
            height_inch = img_h / dpi_y
            img.close()

            if progress_callback:
                progress_callback(60, "正在创建 Word 文档...")

            doc = Document()

            # 页面设置
            section = doc.sections[0]
            page_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches

            # 图片宽度限制在页面宽度内
            if width_inch > page_width:
                scale = page_width / width_inch
                width_inch *= scale
                height_inch *= scale

            # 限制最大高度不超过 9 英寸
            if height_inch > 9.0:
                scale = 9.0 / height_inch
                width_inch *= scale
                height_inch *= scale

            doc.add_picture(img_path, width=Inches(width_inch), height=Inches(height_inch))

            # 添加居中对齐（默认左对齐也可）
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = 1  # 居中

            doc.save(out_path)

            if progress_callback:
                progress_callback(100, "转换完成")

            return {"success": True, "output_files": [out_path], "error": None}
        except Exception as e:
            return {"success": False, "output_files": [], "error": f"JPEG→Word 转换失败：{e}"}

    # ── 调度入口 ──────────────────────────────────────────────

    def convert(
        self,
        input_path: str,
        source_fmt: str,
        target_fmt: str,
        output_dir: str,
        conflict: str = "rename",
        dpi: int = 200,
        quality: int = 90,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None,
        progress_callback: Optional[Callable] = None,
    ) -> dict:
        """
        统一转换调度入口。

        参数:
            input_path:         输入文件的完整路径
            source_fmt:         源格式 ('PDF' / 'Word' / 'JPEG' / 'PNG')
            target_fmt:         目标格式 ('PDF' / 'Word' / 'JPEG')
            output_dir:         输出目录
            conflict:           命名冲突策略 ('rename' / 'overwrite' / 'skip')
            dpi:                渲染 DPI（PDF→JPEG / Word→JPEG）
            quality:            JPEG 质量 (1-100)
            start_page:         起始页（1-based，仅 PDF 源生效）
            end_page:           结束页（1-based，仅 PDF 源生效）
            progress_callback:  callback(percent: int, message: str)

        返回:
            {'success': bool, 'output_files': list[str], 'error': str|None}
        """
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)

        # 规范化格式名
        source = source_fmt.lower()
        target = target_fmt.lower()

        # ── .doc 桥接：先转为临时 .docx ──
        tmp_docx = None
        original_name = None  # 用于传递原始文件名给 word_to_*
        app = None  # COM app 复用
        own_app = False
        if source in ("word", "docx") and input_path.lower().endswith(".doc"):
            if progress_callback:
                progress_callback(0, "检测到 .doc 旧格式，正在转换...")
            original_name = os.path.splitext(os.path.basename(input_path))[0]
            app, _ = self._create_com_app()
            own_app = True
            input_path, tmp_docx = self._ensure_docx(input_path, app, progress_callback)
            source = "docx"  # 后续统一走 docx 路径

        try:
            result = self._convert_inner(
                input_path, source, target, output_dir,
                conflict, dpi, quality, start_page, end_page,
                progress_callback, original_name=original_name, app=app,
            )
        finally:
            if own_app and app is not None:
                try:
                    app.Quit()
                except Exception:
                    pass
            if tmp_docx and os.path.exists(tmp_docx):
                try:
                    os.unlink(tmp_docx)
                except OSError:
                    pass

        return result

    def _convert_inner(
        self,
        input_path: str,
        source: str,
        target: str,
        output_dir: str,
        conflict: str,
        dpi: int,
        quality: int,
        start_page,
        end_page,
        progress_callback,
        original_name: Optional[str] = None,
        app=None,
    ) -> dict:
        """内部转换路由（假设输入已是 .docx 等可直接处理的格式）。"""
        # 路由表
        if source in ("pdf",) and target in ("word", "docx"):
            return self.pdf_to_word(input_path, output_dir, conflict, progress_callback)

        if source in ("pdf",) and target in ("jpeg", "jpg", "png"):
            return self.pdf_to_jpeg(
                input_path, output_dir, dpi, quality,
                start_page, end_page, conflict, progress_callback,
            )

        if source in ("word", "docx") and target == "pdf":
            return self.word_to_pdf(
                input_path, output_dir, conflict,
                original_name=original_name, app=app, progress_callback=progress_callback,
            )

        if source in ("word", "docx") and target in ("jpeg", "jpg", "png"):
            return self.word_to_jpeg(
                input_path, output_dir, dpi, quality, conflict,
                original_name=original_name, app=app, progress_callback=progress_callback,
            )

        if source in ("jpeg", "jpg", "png") and target == "pdf":
            return self.jpeg_to_pdf(input_path, output_dir, conflict, progress_callback)

        if source in ("jpeg", "jpg", "png") and target in ("word", "docx"):
            return self.jpeg_to_word(input_path, output_dir, conflict, progress_callback)

        return {
            "success": False,
            "output_files": [],
            "error": f"不支持的转换路径：{source_fmt} → {target_fmt}",
        }
